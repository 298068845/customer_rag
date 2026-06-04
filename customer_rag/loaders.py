from __future__ import annotations

import tempfile
import zipfile
import posixpath
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable
from xml.etree import ElementTree as ET

from customer_rag.attributes import extract_attributes
from customer_rag.order_flow_image import create_order_flow_image


SUPPORTED_SUFFIXES = {".txt", ".md", ".csv", ".xlsx", ".xls", ".docx", ".pdf"}
SHEET_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
MAIN_DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    source: str
    title: str
    location: str
    image_paths: list[str] | None = None
    tags: list[str] | None = None
    attributes: dict[str, Any] | None = None


LoadDocumentsProgressCallback = Callable[[int, int, Path], None]


def load_documents(
    raw_data_dir: Path,
    progress_callback: LoadDocumentsProgressCallback | None = None,
    source_tags: dict[str, list[str]] | None = None,
) -> list[LoadedDocument]:
    raw_data_dir.mkdir(parents=True, exist_ok=True)
    docs: list[LoadedDocument] = []
    paths = [
        path
        for path in sorted(raw_data_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
    ]
    total = len(paths)
    for index, path in enumerate(paths, start=1):
        docs.extend(load_document_file(path, tags=_tags_for_source(path, source_tags or {})))
        if progress_callback:
            progress_callback(index, total, path)
    return [doc for doc in docs if doc.text.strip()]


def load_document_file(path: Path, tags: list[str] | None = None) -> list[LoadedDocument]:
    if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
        return []
    try:
        initial_tags = _clean_tags(tags or [])
        return [_with_tags(doc, initial_tags) for doc in _load_file(path) if doc.text.strip()]
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"解析文件失败：{path.name}。原因：{exc}") from exc


def _load_file(path: Path) -> Iterable[LoadedDocument]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        yield LoadedDocument(
            text=text,
            source=str(path),
            title=path.name,
            location=path.name,
            attributes=extract_attributes(text),
        )
    elif suffix == ".docx":
        yield from _load_docx(path)
    elif suffix == ".pdf":
        yield from _load_pdf(path)
    elif suffix == ".csv":
        yield from _load_csv(path)
    elif suffix in {".xlsx", ".xls"}:
        yield from _load_excel(path)


def _load_docx(path: Path) -> Iterable[LoadedDocument]:
    try:
        from docx import Document as DocxDocument
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 docx 解析依赖，请先运行：pip install python-docx") from exc

    document = DocxDocument(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(f"表{table_index} 行{row_index}: " + " | ".join(values))
    text = "\n".join(parts)
    yield LoadedDocument(
        text=text,
        source=str(path),
        title=path.name,
        location=path.name,
        attributes=extract_attributes(text),
    )


def _load_pdf(path: Path) -> Iterable[LoadedDocument]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 PDF 解析依赖，请先运行：pip install pypdf") from exc

    reader = PdfReader(str(path))
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            yield LoadedDocument(
                text=text,
                source=str(path),
                title=path.name,
                location=f"{path.name} 第 {page_index} 页",
                attributes=extract_attributes(text),
            )


def _load_csv(path: Path) -> Iterable[LoadedDocument]:
    df = _read_csv_with_fallback(path)
    yield from _rows_to_documents(path, df, path.name)


def _load_excel(path: Path) -> Iterable[LoadedDocument]:
    if path.suffix.lower() == ".xlsx":
        docs = _load_xlsx_package(path)
        if docs:
            yield from docs
            return

    pd = _get_pandas()
    sheets = _read_excel_with_fallback(pd, path)
    for sheet_name, df in sheets.items():
        yield from _rows_to_documents(path, df, f"{path.name} / {sheet_name}")


def _rows_to_documents(path: Path, df: Any, title: str) -> Iterable[LoadedDocument]:
    df = df.fillna("")
    headers = [str(column) for column in df.columns]
    for row_number, row in enumerate(df.itertuples(index=False), start=2):
        fields: dict[str, str] = {}
        for key, value in zip(headers, row):
            field_name = _allowed_field_name(key)
            if not field_name:
                continue
            value_text = str(value).strip()
            if value_text:
                fields[field_name] = value_text
        fields = _prepare_row_fields(fields)
        if _has_core_product_fields(fields):
            text = _fields_to_text(fields)
            image_paths = _order_flow_image_paths(path, row_number, fields)
            yield LoadedDocument(
                text=text,
                source=str(path),
                title=title,
                location=f"{title} 行 {row_number}",
                image_paths=image_paths,
                attributes=extract_attributes(text),
            )


def _load_xlsx_package(path: Path) -> list[LoadedDocument]:
    docs: list[LoadedDocument] = []
    with zipfile.ZipFile(path, "r") as archive:
        shared_strings = _read_shared_strings(archive)
        sheet_paths = _read_workbook_sheets(archive)
        if not sheet_paths:
            return []

        image_map = _extract_xlsx_images(archive, path)
        for sheet_name, sheet_path in sheet_paths:
            rows = _read_sheet_rows(archive, sheet_path, shared_strings)
            if not rows:
                continue

            headers = _pick_headers(rows)
            carry_values: dict[int, str] = {}
            for row_number in sorted(rows):
                if row_number <= 1:
                    continue
                row = rows[row_number]
                fields: dict[str, str] = {}
                for col_index in sorted(headers):
                    header = headers[col_index]
                    field_name = _allowed_field_name(header)
                    if not field_name:
                        continue
                    value = row.get(col_index, "").strip()
                    if not value and _should_fill_down(header):
                        value = carry_values.get(col_index, "")
                    if value:
                        fields[field_name] = value
                        if _should_fill_down(header):
                            carry_values[col_index] = value

                images = image_map.get((sheet_path, row_number), [])
                fields = _prepare_row_fields(fields)

                if _has_core_product_fields(fields):
                    text = _fields_to_text(fields)
                    order_flow_images = _order_flow_image_paths(path, row_number, fields)
                    title = _build_row_title(path.name, sheet_name, row, headers)
                    docs.append(
                        LoadedDocument(
                            text=text,
                            source=str(path),
                            title=title,
                            location=f"{path.name} / {sheet_name} 行 {row_number}",
                            image_paths=order_flow_images + images,
                            attributes=extract_attributes(text),
                        )
                    )
    return docs


def _with_tags(doc: LoadedDocument, initial_tags: list[str]) -> LoadedDocument:
    return LoadedDocument(
        text=doc.text,
        source=doc.source,
        title=doc.title,
        location=doc.location,
        image_paths=doc.image_paths,
        attributes=doc.attributes or extract_attributes(doc.text),
        tags=_clean_tags(
            initial_tags
            + (doc.tags or [])
            + brand_tags_from_text(doc.text)
            + category_tags_from_text(doc.text)
        ),
    )


def brand_tags_from_text(text: str) -> list[str]:
    tags: list[str] = []
    for match in re.finditer(r"(?:^|[；;\n])\s*品牌\s*[:：]\s*([^；;\n]+)", text):
        value = match.group(1).strip()
        if value:
            tags.extend(_split_tag_value(value))
    return _clean_tags(tags)


def category_tags_from_text(text: str) -> list[str]:
    tags: list[str] = []
    for key in ("品类", "类别", "分类"):
        for match in re.finditer(rf"(?:^|[；;\n])\s*{re.escape(key)}\s*[:：]\s*([^；;\n]+)", text):
            value = match.group(1).strip()
            if value:
                tags.extend(_split_tag_value(value))
    return _clean_tags(tags)


def _split_tag_value(value: str) -> list[str]:
    cleaned = re.sub(r"https?://\S+", "", value)
    parts = re.split(r"[,，;；|、/\s]+", cleaned)
    tags: list[str] = []
    for part in parts:
        tag = part.strip(" ：:()（）[]【】")
        if _is_valid_tag_value(tag):
            tags.append(tag)
    return tags


def _is_valid_tag_value(tag: str) -> bool:
    if not (1 <= len(tag) <= 24):
        return False
    if re.fullmatch(r"[-+]?\d+(?:\.\d+)?", tag):
        return False
    blocked_keywords = ("链接", "商品ID", "请勿下单", "隐藏", "http")
    return not any(keyword in tag for keyword in blocked_keywords)


def _tags_for_source(path: Path, source_tags: dict[str, list[str]]) -> list[str]:
    candidates = (
        str(path),
        str(path.resolve()),
        path.name,
        path.stem,
    )
    for candidate in candidates:
        tags = source_tags.get(candidate)
        if tags:
            return tags
    return []


def _clean_tags(tags: list[str]) -> list[str]:
    cleaned: list[str] = []
    for tag in tags:
        value = str(tag).strip()
        if value and value not in cleaned:
            cleaned.append(value)
    return cleaned


def _read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    strings: list[str] = []
    for item in root.findall(f"{{{SHEET_NS}}}si"):
        text = "".join(node.text or "" for node in item.iter(f"{{{SHEET_NS}}}t"))
        strings.append(text)
    return strings


def _read_workbook_sheets(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    if "xl/workbook.xml" not in archive.namelist():
        return []

    rels = _read_relationships(archive, "xl/_rels/workbook.xml.rels")
    root = ET.fromstring(archive.read("xl/workbook.xml"))
    sheets: list[tuple[str, str]] = []
    for sheet in root.findall(f".//{{{SHEET_NS}}}sheet"):
        name = sheet.attrib.get("name", "Sheet")
        rel_id = sheet.attrib.get(f"{{{OFFICE_REL_NS}}}id")
        target = rels.get(rel_id or "")
        if not target:
            continue
        sheet_path = _resolve_package_path("xl/workbook.xml", target)
        sheets.append((name, sheet_path))
    return sheets


def _read_sheet_rows(
    archive: zipfile.ZipFile,
    sheet_path: str,
    shared_strings: list[str],
) -> dict[int, dict[int, str]]:
    root = ET.fromstring(archive.read(sheet_path))
    rows: dict[int, dict[int, str]] = {}
    for row in root.findall(f".//{{{SHEET_NS}}}row"):
        row_number = int(row.attrib.get("r", "0"))
        values: dict[int, str] = {}
        for cell in row.findall(f"{{{SHEET_NS}}}c"):
            ref = cell.attrib.get("r", "")
            col_index = _column_index(ref)
            if col_index < 1:
                continue
            values[col_index] = _read_cell_value(cell, shared_strings)
        rows[row_number] = values
    return rows


def _read_cell_value(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.iter(f"{{{SHEET_NS}}}t")).strip()

    value_node = cell.find(f"{{{SHEET_NS}}}v")
    if value_node is None or value_node.text is None:
        return ""

    value = value_node.text
    if cell_type == "s":
        try:
            return shared_strings[int(value)].strip()
        except (IndexError, ValueError):
            return value.strip()
    return value.strip()


def _pick_headers(rows: dict[int, dict[int, str]]) -> dict[int, str]:
    first_row = rows.get(1, {})
    headers = {col: value.replace("\n", " ").strip() for col, value in first_row.items() if value.strip()}
    if headers:
        return headers
    max_cols = max((max(row.keys(), default=0) for row in rows.values()), default=0)
    return {col: f"列{col}" for col in range(1, max_cols + 1)}


def _extract_xlsx_images(archive: zipfile.ZipFile, workbook_path: Path) -> dict[tuple[str, int], list[str]]:
    image_map: dict[tuple[str, int], list[str]] = {}
    sheet_paths = {path for _, path in _read_workbook_sheets(archive)}
    assets_dir = workbook_path.parent / "_assets" / workbook_path.stem
    names = set(archive.namelist())

    for sheet_path in sheet_paths:
        sheet_rels_path = _rels_path(sheet_path)
        sheet_rels = _read_relationships(archive, sheet_rels_path)
        if not sheet_rels:
            continue

        for target in sheet_rels.values():
            drawing_path = _resolve_package_path(sheet_path, target)
            if drawing_path not in names or not drawing_path.startswith("xl/drawings/"):
                continue
            drawing_rels = _read_relationships(archive, _rels_path(drawing_path))
            drawing_root = ET.fromstring(archive.read(drawing_path))
            for anchor in list(drawing_root):
                from_node = anchor.find(f"{{{DRAWING_NS}}}from")
                blip = anchor.find(f".//{{{MAIN_DRAWING_NS}}}blip")
                if from_node is None or blip is None:
                    continue
                row_node = from_node.find(f"{{{DRAWING_NS}}}row")
                rel_id = blip.attrib.get(f"{{{OFFICE_REL_NS}}}embed") or blip.attrib.get(f"{{{OFFICE_REL_NS}}}link")
                if row_node is None or not rel_id:
                    continue
                media_target = drawing_rels.get(rel_id)
                if not media_target:
                    continue
                media_path = _resolve_package_path(drawing_path, media_target)
                if media_path not in names:
                    continue

                assets_dir.mkdir(parents=True, exist_ok=True)
                output_path = assets_dir / Path(media_path).name
                if not output_path.exists():
                    output_path.write_bytes(archive.read(media_path))
                row_number = int(row_node.text or "0") + 1
                image_map.setdefault((sheet_path, row_number), []).append(str(output_path))
    return image_map


def _read_relationships(archive: zipfile.ZipFile, rels_path: str) -> dict[str, str]:
    if rels_path not in archive.namelist():
        return {}
    root = ET.fromstring(archive.read(rels_path))
    return {
        rel.attrib["Id"]: rel.attrib["Target"]
        for rel in root.findall(f"{{{REL_NS}}}Relationship")
        if rel.attrib.get("Id") and rel.attrib.get("Target")
    }


def _rels_path(part_path: str) -> str:
    part = Path(part_path)
    return str(part.parent / "_rels" / f"{part.name}.rels").replace("\\", "/")


def _resolve_package_path(base_part_path: str, target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    base_dir = Path(base_part_path).parent
    return posixpath.normpath((base_dir / target).as_posix())


def _column_index(cell_ref: str) -> int:
    letters = ""
    for char in cell_ref:
        if char.isalpha():
            letters += char.upper()
        else:
            break
    index = 0
    for char in letters:
        index = index * 26 + ord(char) - 64
    return index


def _should_fill_down(header: str) -> bool:
    return _allowed_field_name(header) in {"品牌", "品类", "商品链接", "其他说明"}


def _order_flow_image_paths(path: Path, row_number: int, fields: dict[str, str]) -> list[str]:
    image_path = create_order_flow_image(
        fields.get("产品信息", ""),
        fields.get("下单流程", ""),
        fields.get("其他说明", ""),
        source_path=path,
        row_number=row_number,
        output_root=path.parent / "_generated" / "order_flow",
    )
    return [image_path] if image_path else []


def _allowed_field_name(header: str) -> str | None:
    normalized = re.sub(r"\s+", "", str(header)).lower()
    if not normalized:
        return None
    if "商品id" in normalized or "skuid" in normalized or "skuid" in normalized.replace("/", ""):
        return None
    if "图片" in normalized or "图片" in str(header):
        return None
    if "品牌" in normalized:
        return "品牌"
    if "品类" in normalized or "类别" in normalized or "分类" in normalized:
        return "品类"
    if "产品信息" in normalized or "型号" in normalized or "规格" in normalized:
        return "产品信息"
    if "下单流程" in normalized or "付款流程" in normalized:
        return "下单流程"
    if "商品链接" in normalized or "礼金短链接" in normalized or normalized == "链接":
        return "商品链接"
    if "其他说明" in normalized or "特殊信息" in normalized or "限制" in normalized:
        return "其他说明"
    if "其他链接" in normalized:
        return "其他链接"
    return None


def _has_core_product_fields(fields: dict[str, str]) -> bool:
    field_names = {field for field, value in fields.items() if str(value).strip()}
    core_fields = {"品牌", "品类", "产品信息", "下单流程", "其他说明"}
    return bool(field_names & core_fields)


def _prepare_row_fields(fields: dict[str, str]) -> dict[str, str]:
    prepared = {key: str(value).strip() for key, value in fields.items() if str(value).strip()}
    extracted_links: list[str] = []
    for field_name in ("产品信息", "下单流程", "其他说明"):
        value = prepared.get(field_name, "")
        if not value:
            continue
        cleaned, links = _extract_links(value)
        prepared[field_name] = cleaned
        extracted_links.extend(links)
    if extracted_links:
        prepared["其他链接"] = _join_unique_links([prepared.get("其他链接", ""), *extracted_links])
    return {key: value for key, value in prepared.items() if value}


def _fields_to_text(fields: dict[str, str]) -> str:
    order = ("品牌", "品类", "产品信息", "下单流程", "商品链接", "其他说明", "其他链接")
    pairs = [f"{field}: {fields[field]}" for field in order if fields.get(field)]
    pairs.extend(f"{field}: {value}" for field, value in fields.items() if field not in order and value)
    return "；".join(pairs)


def _extract_links(text: str) -> tuple[str, list[str]]:
    links = re.findall(r"https?://[^\s，,；;。)）】]+", text)
    if not links:
        return text.strip(), []
    cleaned = text
    for link in links:
        cleaned = cleaned.replace(link, "")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\s*([，,；;。])\s*", r"\1", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip(" ，,；;。")
    return cleaned.strip(), links


def _join_unique_links(values: list[str]) -> str:
    links: list[str] = []
    for value in values:
        for link in re.findall(r"https?://[^\s，,；;。)）】]+", str(value)):
            if link not in links:
                links.append(link)
    return "\n".join(links)


def _build_row_title(file_name: str, sheet_name: str, row: dict[int, str], headers: dict[int, str]) -> str:
    for keyword in ("产品信息", "商品", "名称", "型号"):
        for col_index, header in headers.items():
            if keyword in header and row.get(col_index, "").strip():
                return row[col_index].strip().splitlines()[0][:80]
    return f"{file_name} / {sheet_name}"


def _read_csv_with_fallback(path: Path) -> pd.DataFrame:
    pd = _get_pandas()
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return pd.read_csv(path, encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
    if last_error:
        raise last_error
    return pd.read_csv(path)


def _get_pandas() -> Any:
    try:
        import pandas as pd
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少表格解析依赖，请先运行：pip install pandas openpyxl xlrd") from exc
    return pd


def _read_excel_with_fallback(pd: Any, path: Path) -> dict[str, Any]:
    try:
        return pd.read_excel(path, sheet_name=None)
    except Exception as exc:
        if path.suffix.lower() != ".xlsx":
            raise
        try:
            return _read_xlsx_without_styles(pd, path)
        except Exception as fallback_exc:
            raise RuntimeError(
                f"Excel 文件读取失败：{path.name}。如果这是腾讯文档导出的表格，请尝试在腾讯文档中重新导出为 CSV 后导入；"
                f"原始错误：{exc}；兜底读取错误：{fallback_exc}"
            ) from fallback_exc


def _read_xlsx_without_styles(pd: Any, path: Path) -> dict[str, Any]:
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
            for info in src.infolist():
                if info.filename == "xl/styles.xml":
                    continue
                dst.writestr(info, src.read(info.filename))
        return pd.read_excel(tmp_path, sheet_name=None)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()
